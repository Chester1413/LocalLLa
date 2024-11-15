import fitz  # PyMuPDF
import requests
from docx import Document as DocxDocument
import textract
import os

def read_file(file_path):
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return read_pdf(file_path)
        elif ext == ".docx":
            return read_docx(file_path)
        elif ext == ".txt":
            return read_txt(file_path)
        elif ext == ".msg":
            return read_msg(file_path)
        else:
            print(f"不支持的文件類型：{ext}")
            return ""
    except Exception as e:
        print(f"無法讀取文件：{e}")
        return ""

def read_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"無法讀取PDF文件：{e}")
        return ""

def read_docx(docx_path):
    try:
        doc = DocxDocument(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"無法讀取DOCX文件：{e}")
        return ""

def read_txt(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            print(txt_path)
            return file.read()
    except Exception as e:
        print(f"無法讀取TXT文件：{e}")
        return ""

def read_msg(msg_path):
    try:
        text = textract.process(msg_path)
        return text.decode('utf-8')
    except Exception as e:
        print(f"無法讀取MSG文件：{e}")
        return ""

def get_user_input(first_time=True):
    if first_time:
        choice = input("請選擇：輸入資料夾位置 輸入1，跳過輸入2：").strip()
        if choice == "1":
            folder_path = input("請輸入資料夾位置：").strip()
        else:
            folder_path = None
    else:
        folder_path = None
    prompt = input("請輸入prompt：")
    return folder_path, prompt

def read_all_files_in_folder(folder_path):
    all_text = ""
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_text = read_file(file_path)
            if file_text:
                all_text += f"\n\n文件：{file}\n{file_text}"
    return all_text

def send_request_to_oobabooga(folder_path, prompt, history):
    url = "http://127.0.0.1:5000/v1/chat/completions"
    headers = {
        "Content-Type": "application/json"
    }

    if folder_path:
        folder_text = read_all_files_in_folder(folder_path)
        if folder_text:
            prompt = f"{prompt}\n\n資料夾內容：\n{folder_text}"
        else:
            print("資料夾內沒有可讀取的文件或讀取失敗。")
            return
    else:
        print("")

    history.append({"role": "user", "content": prompt})
    data = {
        "mode": "chat-instruct",
        "character": "Example",
        "messages": history
    }

    try:
        response = requests.post(url, headers=headers, json=data, verify=False)
        response.raise_for_status()  # 檢查請求是否成功
        assistant_message = response.json()['choices'][0]['message']['content']
        history.append({"role": "assistant", "content": assistant_message})
        print(assistant_message)
    except requests.exceptions.RequestException as e:
        print(f"請求失敗：{e}")

def save_conversation_to_docx(history, filename="conversation.docx"):
    doc = DocxDocument()
    for message in history:
        if message["role"] == "user":
            # 只保存使用者的原始 prompt，不包括資料夾內容
            original_prompt = message['content'].split("\n\n資料夾內容：")[0]
            doc.add_paragraph(f"User: {original_prompt}")
        elif message["role"] == "assistant":
            doc.add_paragraph(f"Assistant: {message['content']}")
    doc.save(filename)
    print(f"對話已保存到 {filename}")



    def main():
        history = []
        first_time = True
        while True:
            folder_path, prompt = get_user_input(first_time)
            first_time = False
            if prompt.lower() == "exit":
                print("退出對話。")
                break
            elif prompt.lower() == "exit2":
                save_conversation_to_docx(history)
                print("退出並保存對話。")
                break
            send_request_to_oobabooga(folder_path, prompt, history)

    if __name__ == "__main__":
        main()
